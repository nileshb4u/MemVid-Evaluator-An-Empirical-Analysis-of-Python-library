# memvid-evaluator/preprocessor.py

import os
import magic # python-magic
from io import BytesIO

# For PDF processing
from PyPDF2 import PdfReader

# For DOCX processing
from docx import Document as DocxDocument # Renamed to avoid conflict with any Document class you might create

# Import configuration if needed, though not strictly for parsing functions themselves
# import config

# Supported file types and their magic identifiers (these might need refinement)
# You can find magic numbers/strings by inspecting files with `file` command on Linux/macOS
# or using python-magic to inspect known file types.
# For simplicity, we'll primarily rely on extensions but use magic as a fallback/verification.

SUPPORTED_EXTENSIONS = {
    ".txt": "text/plain",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

def get_file_type(file_path_or_buffer):
    """
    Determines the file type using python-magic and file extension.
    :param file_path_or_buffer: Path to the file or a file-like object (BytesIO).
    :return: A tuple (detected_mime_type, file_extension) or (None, None) if undetectable.
    """
    detected_mime_type = None
    file_extension = None

    try:
        if isinstance(file_path_or_buffer, str): # It's a file path
            file_extension = os.path.splitext(file_path_or_buffer)[1].lower()
            # For local files, magic can read directly from path
            detected_mime_type = magic.from_file(file_path_or_buffer, mime=True)
        elif hasattr(file_path_or_buffer, 'read') and hasattr(file_path_or_buffer, 'seek'): # It's a file-like object
            # For BytesIO or uploaded files, read a small chunk for magic, then reset pointer
            initial_pos = file_path_or_buffer.tell()
            file_sample = file_path_or_buffer.read(2048) # Read first 2KB for MIME type detection
            file_path_or_buffer.seek(initial_pos) # Reset buffer to original position
            detected_mime_type = magic.from_buffer(file_sample, mime=True)

            # Try to get extension from name attribute if available (e.g., Streamlit UploadedFile)
            if hasattr(file_path_or_buffer, 'name'):
                file_extension = os.path.splitext(file_path_or_buffer.name)[1].lower()
        else:
            print("Error: Input is not a valid file path or buffer.")
            return None, None

    except Exception as e:
        print(f"Error detecting file type: {e}")
        # Fallback to extension if magic fails but we have a path
        if isinstance(file_path_or_buffer, str) and not file_extension:
            file_extension = os.path.splitext(file_path_or_buffer)[1].lower()
        return None, file_extension # Return None for MIME if magic failed

    return detected_mime_type, file_extension


def _parse_txt(file_buffer):
    """Parses a .txt file."""
    try:
        # Assuming UTF-8, but could add more robust encoding detection if needed
        return file_buffer.read().decode('utf-8')
    except Exception as e:
        print(f"Error parsing TXT file: {e}")
        # Try with 'latin-1' as a fallback for some common non-UTF-8 files
        try:
            file_buffer.seek(0) # Reset buffer
            return file_buffer.read().decode('latin-1')
        except Exception as e_fallback:
            print(f"Fallback TXT parsing also failed: {e_fallback}")
            return None

def _parse_pdf(file_buffer):
    """Parses a .pdf file using PyPDF2."""
    text_content = ""
    try:
        reader = PdfReader(file_buffer)
        for page in reader.pages:
            text = page.extract_text()
            if text: # Ensure text is not None
                text_content += text + "\n" # Add newline between pages
        return text_content.strip() if text_content else None
    except Exception as e:
        print(f"Error parsing PDF file with PyPDF2: {e}")
        return None

def _parse_docx(file_buffer):
    """Parses a .docx file using python-docx."""
    text_content = ""
    try:
        document = DocxDocument(file_buffer)
        for para in document.paragraphs:
            text_content += para.text + "\n"
        return text_content.strip() if text_content else None
    except Exception as e:
        print(f"Error parsing DOCX file: {e}")
        return None

def extract_text_from_file(file_path_or_buffer):
    """
    Extracts text content from a given file (path or buffer).
    Uses file extension primarily and mime type as a secondary check.

    :param file_path_or_buffer: Path to the file (str) or a file-like object (e.g., BytesIO, Streamlit UploadedFile).
    :return: Extracted text as a string, or None if extraction fails or type is unsupported.
    """
    mime_type, extension = get_file_type(file_path_or_buffer)

    # If input is a path, open it as a BytesIO buffer for consistent handling
    # If it's already a buffer, ensure it's in binary mode (Streamlit UploadedFile usually is)
    file_buffer = None
    is_path = isinstance(file_path_or_buffer, str)

    try:
        if is_path:
            with open(file_path_or_buffer, "rb") as f:
                file_buffer = BytesIO(f.read())
        elif hasattr(file_path_or_buffer, 'read') and hasattr(file_path_or_buffer, 'seek'):
            # Ensure we are at the beginning of the buffer if it's already opened
            file_path_or_buffer.seek(0)
            file_buffer = file_path_or_buffer
        else:
            print("Error: Invalid input type for file processing.")
            return None

        if not file_buffer:
            return None # Should not happen if logic above is correct

        # Primary dispatch based on extension
        if extension == ".txt":
            return _parse_txt(file_buffer)
        elif extension == ".pdf":
            # Mime check for PDF can be useful as .pdf files might not always have text
            if mime_type == "application/pdf":
                return _parse_pdf(file_buffer)
            else:
                print(f"Warning: File with .pdf extension is not identified as PDF by magic: {mime_type}")
                # Attempt parsing anyway if extension is .pdf
                return _parse_pdf(file_buffer)
        elif extension == ".docx":
            # Check against known DOCX MIME type
            if mime_type == SUPPORTED_EXTENSIONS[".docx"]:
                 return _parse_docx(file_buffer)
            else:
                print(f"Warning: File with .docx extension is not identified as DOCX by magic: {mime_type}")
                # Attempt parsing anyway if extension is .docx
                return _parse_docx(file_buffer)
        else:
            print(f"Unsupported file extension: {extension} (MIME: {mime_type})")
            return None

    except Exception as e:
        print(f"An unexpected error occurred during text extraction: {e}")
        return None
    finally:
        # If we opened the file (from path), BytesIO doesn't need explicit closing here.
        # If the buffer was passed in, the caller is responsible for it.
        pass


# --- Example Usage (for testing this module directly) ---
if __name__ == '__main__':
    # Create dummy files for testing
    # Ensure you have a 'data/input_docs' directory or adjust path
    sample_docs_dir = os.path.join(os.path.dirname(__file__), 'data', 'input_docs')
    os.makedirs(sample_docs_dir, exist_ok=True)

    # Create a dummy TXT file
    txt_file_path = os.path.join(sample_docs_dir, "sample.txt")
    with open(txt_file_path, "w", encoding="utf-8") as f:
        f.write("This is a sample text file.\nIt has multiple lines.\nHello, World!")
    print(f"\n--- Testing TXT file: {txt_file_path} ---")
    text_from_txt = extract_text_from_file(txt_file_path)
    if text_from_txt:
        print("Extracted Text (TXT):\n", text_from_txt)
    else:
        print("Failed to extract text from TXT.")

    # Create a dummy DOCX file (requires python-docx to create, or use a pre-existing one)
    # For simplicity, this example assumes you might place a sample.docx there.
    # If you want to create one programmatically:
    # from docx import Document as CreateDocx
    # doc = CreateDocx()
    # doc.add_paragraph("This is a sample DOCX document created for testing.")
    # doc.add_paragraph("It also has multiple paragraphs.")
    # docx_file_path = os.path.join(sample_docs_dir, "sample.docx")
    # doc.save(docx_file_path)

    docx_file_path = os.path.join(sample_docs_dir, "sample.docx") # Assume you place one here
    if os.path.exists(docx_file_path):
        print(f"\n--- Testing DOCX file: {docx_file_path} ---")
        text_from_docx = extract_text_from_file(docx_file_path)
        if text_from_docx:
            print("Extracted Text (DOCX):\n", text_from_docx)
        else:
            print("Failed to extract text from DOCX.")
    else:
        print(f"\nSkipping DOCX test: Place a 'sample.docx' in '{sample_docs_dir}'")


    # Create a dummy PDF file (requires a library like reportlab to create, or use a pre-existing one)
    # For simplicity, this example assumes you might place a sample.pdf there.
    # If you want to create one programmatically:
    # from reportlab.pdfgen import canvas
    # pdf_file_path = os.path.join(sample_docs_dir, "sample.pdf")
    # c = canvas.Canvas(pdf_file_path)
    # c.drawString(100, 750, "This is a sample PDF document.")
    # c.drawString(100, 730, "Created for testing purposes.")
    # c.save()
    pdf_file_path = os.path.join(sample_docs_dir, "sample.pdf") # Assume you place one here
    if os.path.exists(pdf_file_path):
        print(f"\n--- Testing PDF file: {pdf_file_path} ---")
        text_from_pdf = extract_text_from_file(pdf_file_path)
        if text_from_pdf:
            print("Extracted Text (PDF):\n", text_from_pdf)
        else:
            print("Failed to extract text from PDF.")
    else:
        print(f"\nSkipping PDF test: Place a 'sample.pdf' in '{sample_docs_dir}'")

    # Test with a non-existent file
    print("\n--- Testing non-existent file ---")
    text_non_existent = extract_text_from_file("non_existent_file.xyz")
    if not text_non_existent:
        print("Correctly handled non-existent file.")

    # Test with an unsupported file type (e.g., an image)
    # Create a dummy PNG file (or use any image)
    # For simplicity, this example assumes you might place a sample.png there.
    # from PIL import Image
    # img = Image.new('RGB', (60, 30), color = 'red')
    # png_file_path = os.path.join(sample_docs_dir, "sample.png")
    # img.save(png_file_path)
    png_file_path = os.path.join(sample_docs_dir, "sample.png") # Assume you place one here
    if os.path.exists(png_file_path):
        print(f"\n--- Testing unsupported file (PNG): {png_file_path} ---")
        text_from_png = extract_text_from_file(png_file_path)
        if not text_from_png:
            print("Correctly handled unsupported PNG file.")
    else:
        print(f"\nSkipping PNG test: Place a 'sample.png' in '{sample_docs_dir}'")